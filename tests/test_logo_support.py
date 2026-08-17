from __future__ import annotations

from pathlib import Path

from PIL import Image
from docx import Document

from document_automation_studio.engine.preview_engine import PreviewEngine
from document_automation_studio.engine.rule_engine import RuleEngine
from document_automation_studio.models.rule_models import RuleSet, TextReplacementRule
from document_automation_studio.processors.word_processor import WordProcessor


def _create_logo_file(path: Path) -> None:
    image = Image.new("RGB", (100, 100), color=(255, 0, 0))
    image.save(path)


def test_word_processor_inserts_logo_at_placeholder(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    logo_path = tmp_path / "logo.png"
    _create_logo_file(logo_path)

    document = Document()
    document.add_paragraph("This is a document with a {{Logo}} placeholder.")
    document.save(source)

    processor = WordProcessor()
    output_root = tmp_path / "out"
    output_root.mkdir(exist_ok=True)

    destination = processor.process(
        source_path=source,
        output_root=output_root,
        preserve_folder_structure=False,
        logo_path=str(logo_path),
    )

    assert destination.exists()
    loaded = Document(destination)
    assert any("{{Logo}}" not in paragraph.text for paragraph in loaded.paragraphs)


def test_preview_engine_detects_logo_insertion(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Insert logo here: {{Logo}}")
    document.save(source)

    engine = PreviewEngine(RuleEngine())
    preview_items = engine.preview_files(
        [source],
        rule_set=RuleSet(),
        input_root=tmp_path,
        max_items=10,
        process_word_files=True,
        process_excel_files=False,
        logo_path="logo.png",
    )

    assert any(item.change_type == "Logo Insertion" for item in preview_items)
