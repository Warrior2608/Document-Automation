from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

from document_automation_studio.processors.branding_processor import BrandingProcessor
from document_automation_studio.processors.word_processor import WordProcessor


def test_word_processor_replaces_placeholders(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Welcome {{Company}} to the system.")
    document.save(source)

    processor = WordProcessor()
    output_root = tmp_path / "out"
    output_root.mkdir(exist_ok=True)

    destination = processor.process(
        source_path=source,
        output_root=output_root,
        preserve_folder_structure=False,
        replacements={"{{Company}}": "ACME Corp"},
    )

    assert destination.exists()
    loaded = Document(destination)
    assert any("ACME Corp" in paragraph.text for paragraph in loaded.paragraphs)


def test_word_processor_inserts_logo_with_preserved_ratio(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    placeholder_path = tmp_path / "placeholder.png"
    Image.new("RGBA", (400, 200), (255, 0, 0, 255)).save(placeholder_path)
    run.add_picture(str(placeholder_path), width=Inches(2.0))
    document.save(source)

    logo_path = tmp_path / "logo.png"
    Image.new("RGBA", (400, 200), (255, 0, 0, 255)).save(logo_path)

    processor = WordProcessor()
    output_root = tmp_path / "out"
    output_root.mkdir(exist_ok=True)

    destination = processor.process(
        source_path=source,
        output_root=output_root,
        preserve_folder_structure=False,
        logo_path=str(logo_path),
    )

    assert destination.exists()
    updated = Document(destination)
    updated_run = updated.paragraphs[0].runs[0]
    xml = updated_run._element.xml

    assert "<wp:extent" in xml
    assert "<a:ext" in xml
    assert "cx=\"1828800\"" in xml
    assert "cy=\"914400\"" in xml


def test_branding_replaces_header_text_case_insensitively() -> None:
    document = Document()
    header_paragraph = document.sections[0].header.paragraphs[0]
    header_paragraph.text = "ACME CORP"

    processor = BrandingProcessor()
    replaced = processor._replace_text_in_paragraph(header_paragraph, "Acme Corp", "Contoso")

    assert replaced == 1
    assert header_paragraph.text == "Contoso"


def test_branding_replaces_body_text_case_insensitively() -> None:
    document = Document()
    body_paragraph = document.add_paragraph("This document is owned by ACME CORP.")

    processor = BrandingProcessor()
    replaced = processor._replace_text_in_paragraph(body_paragraph, "Acme Corp", "Contoso")

    assert replaced == 1
    assert body_paragraph.text == "This document is owned by Contoso."
