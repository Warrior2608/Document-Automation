from pathlib import Path
from docx import Document
from document_automation_studio.processors.branding_processor import BrandingProcessor, BrandingSettings

path = Path('C:/temp/header_fix_check.docx')
path.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
section = doc.sections[0]
section.header.paragraphs[0].text = 'ACME CORP'
section.footer.paragraphs[0].text = 'Footer ACME CORP'
doc.save(path)

result = BrandingProcessor().apply_to_document(
    path,
    BrandingSettings(
        current_company_name='ACME CORP',
        new_company_name='Contoso',
        replace_in_header=True,
        replace_in_footer=True,
    ),
)

reloaded = Document(path)
header_text = reloaded.sections[0].header.paragraphs[0].text
footer_text = reloaded.sections[0].footer.paragraphs[0].text
print('result.success=', result.success)
print('result.replacements_made=', result.replacements_made)
print('header=', header_text)
print('footer=', footer_text)
print('PASS' if result.success and header_text == 'Contoso' and footer_text == 'Footer Contoso' else 'FAIL')
