from __future__ import annotations

from pathlib import Path

from docx import Document

from document_automation_studio.processors.branding_processor import BrandingProcessor, BrandingSettings


def test_branding_replacement_updates_header_text(tmp_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "ABC Technologies Sdn. Bhd."

    path = tmp_path / "branding.docx"
    document.save(path)

    processor = BrandingProcessor(learning_db_path=tmp_path / "branding_learning.sqlite")
    settings = BrandingSettings(current_company_name="ABC Technologies Sdn. Bhd.", new_company_name="XYZ Technologies Pvt. Ltd.", replace_in_header=True)
    result = processor.apply_to_document(path, settings)

    updated_document = Document(path)
    updated_header_text = updated_document.sections[0].header.paragraphs[0].text

    assert result.success is True
    assert "XYZ Technologies Pvt. Ltd." in updated_header_text
    assert result.replacements_made >= 1
