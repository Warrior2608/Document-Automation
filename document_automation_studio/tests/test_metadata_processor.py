from __future__ import annotations

from pathlib import Path

from docx import Document

from document_automation_studio.processors.metadata_processor import MetadataProcessor, MetadataValues


def test_metadata_replacement_updates_only_values_and_preserves_labels(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    row = table.rows[0]
    row.cells[0].text = "Effective Date"
    row.cells[1].text = "22-07-2026"

    row2 = table.add_row().cells
    row2[0].text = "Prepared By (Name / Designation)"
    row2[1].text = "Nur Syafirah Mazlan / Director"

    document_path = tmp_path / "sample.docx"
    document.save(document_path)

    processor = MetadataProcessor()
    values = MetadataValues(effective_date="01-08-2026", prepared_by_name="John Smith", prepared_by_designation="Quality Manager")
    result = processor.apply_to_document(document_path, values)

    loaded_document = Document(document_path)
    table = loaded_document.tables[0]

    assert table.rows[0].cells[0].text == "Effective Date"
    assert table.rows[0].cells[1].text == "01-08-2026"
    assert table.rows[1].cells[0].text == "Prepared By (Name / Designation)"
    assert table.rows[1].cells[1].text == "John Smith / Quality Manager"
    assert result.metadata_table_found is True
    assert result.success is True
    assert "effective_date" in result.updated_fields
    assert "prepared_by" in result.updated_fields
