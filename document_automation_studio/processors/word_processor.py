from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Emu, Inches
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from PIL import Image

from document_automation_studio.engine.rule_engine import RuleEngine
from document_automation_studio.models.rule_models import RuleSet, TextReplacementRule
from document_automation_studio.processors.branding_processor import BrandingProcessor, BrandingSettings
from document_automation_studio.processors.metadata_processor import MetadataProcessor, MetadataValues

logger = logging.getLogger(__name__)


class WordProcessor:
    """Processor for Word documents with batch-ready replacement and save behavior."""

    def __init__(self) -> None:
        self.logger = logger
        self.rule_engine = RuleEngine()
        self.metadata_processor = MetadataProcessor()
        self.branding_processor = BrandingProcessor()

    def process(
        self,
        source_path: Path,
        output_root: Path,
        preserve_folder_structure: bool = True,
        input_root: Path | None = None,
        replacements: dict[str, str] | None = None,
        rule_set: RuleSet | None = None,
        logo_path: str | None = None,
        logo_width_in: float | None = None,
        logo_height_in: float | None = None,
        metadata_values: MetadataValues | None = None,
        branding_settings: BrandingSettings | None = None,
    ) -> Path:
        """Process a .docx file and save the result into the output root."""
        if source_path.suffix.lower() != ".docx":
            raise ValueError("Unsupported file type for WordProcessor: %s" % source_path)

        document = Document(source_path)
        self.logger.debug("Processing Word document %s", source_path)

        if replacements:
            replacement_rules = [
                TextReplacementRule(find=find_text, replace=replace_text)
                for find_text, replace_text in replacements.items()
            ]
            self.replace_placeholders(document, replacement_rules)
        elif rule_set and rule_set.text_replacements:
            self.replace_placeholders(document, rule_set.text_replacements)

        if logo_path:
            self._insert_logo(document, logo_path, width_in=logo_width_in, height_in=logo_height_in)

        if metadata_values is not None:
            try:
                self.metadata_processor.apply_to_document_object(document, metadata_values, document_name=source_path.name)
            except Exception as error:  # pragma: no cover - defensive path
                self.logger.warning("Metadata replacement skipped for %s: %s", source_path, error)

        if branding_settings is not None:
            try:
                self.branding_processor.apply_to_document_object(document, branding_settings, document_name=source_path.name)
            except Exception as error:  # pragma: no cover - defensive path
                self.logger.warning("Branding replacement skipped for %s: %s", source_path, error)

        destination = self._build_destination(source_path, output_root, preserve_folder_structure, input_root)
        output_root.mkdir(parents=True, exist_ok=True)
        destination.parent.mkdir(parents=True, exist_ok=True)
        document.save(destination)
        self.logger.info("Word document saved to %s", destination)

        if not destination.exists():
            raise FileNotFoundError(f"Expected output file was not created: {destination}")
        self.logger.info("Word document saved to %s", destination)
        return destination

    def replace_placeholders(self, document: DocxDocument, replacements: Iterable[TextReplacementRule]) -> None:
        """Replace placeholders across paragraphs, tables, headers, and footers."""
        self.logger.debug("Replacing placeholders in document")
        for paragraph in document.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)

        for table in document.tables:
            self._replace_text_in_table(table, replacements)

        for section in document.sections:
            for header in (section.header, section.first_page_header, section.even_page_header):
                if header is not None:
                    for paragraph in header.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)
                    for table in header.tables:
                        self._replace_text_in_table(table, replacements)

            for footer in (section.footer, section.first_page_footer, section.even_page_footer):
                if footer is not None:
                    for paragraph in footer.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)
                    for table in footer.tables:
                        self._replace_text_in_table(table, replacements)

    def _replace_text_in_paragraph(self, paragraph: Paragraph, replacements: Iterable[TextReplacementRule]) -> None:
        text = paragraph.text
        text = self.rule_engine.apply_text_replacements(text, replacements)
        if text != paragraph.text:
            paragraph.text = text

    def _replace_text_in_table(self, table: _Cell | object, replacements: Iterable[TextReplacementRule]) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_text_in_paragraph(paragraph, replacements)

    def _replace_text(self, source: str, find_text: str, replace_text: str) -> str:
        return source.replace(find_text, replace_text)

    def _insert_logo(self, document: DocxDocument, logo_path: str, width_in: float | None = None, height_in: float | None = None) -> None:
        try:
            image_path = Path(logo_path)
            if not image_path.exists():
                self.logger.warning("Logo file not found: %s", logo_path)
                return

            desired_width = int((width_in if width_in is not None else 1.8) * 914400)
            desired_height = int((height_in if height_in is not None else 0.55) * 914400)

            EMU_PER_INCH = 914400

            def _get_run_image_extent(run) -> tuple[int | None, int | None]:
                try:
                    for el in run._element.iter():
                        tag = getattr(el, 'tag', '')
                        if tag and tag.endswith('}extent'):
                            cx = el.get('cx')
                            cy = el.get('cy')
                            return (int(cx) if cx else None, int(cy) if cy else None)
                except Exception:
                    return (None, None)
                return (None, None)

            def _set_extent_on_run(run, width_emus: int | None, height_emus: int | None) -> None:
                try:
                    for el in run._element.iter():
                        tag = getattr(el, 'tag', '')
                        if tag and tag.endswith('}extent'):
                            if width_emus is not None:
                                el.set('cx', str(width_emus))
                            if height_emus is not None:
                                el.set('cy', str(height_emus))
                        if tag and tag.endswith('}ext'):
                            if width_emus is not None:
                                el.set('cx', str(width_emus))
                            if height_emus is not None:
                                el.set('cy', str(height_emus))
                except Exception:
                    pass

            def _remove_drawings_from_run(run) -> bool:
                removed = False
                try:
                    for el in list(run._element.iter()):
                        tag = getattr(el, 'tag', '')
                        if tag and (tag.endswith('}drawing') or tag.endswith('}pict')):
                            parent = el.getparent()
                            if parent is not None:
                                parent.remove(el)
                                removed = True
                except Exception:
                    pass
                return removed

            def _get_image_native_size(path: Path) -> tuple[int, int]:
                with Image.open(path) as img:
                    return img.width, img.height

            def _calculate_fit_size(orig_cx: int | None, orig_cy: int | None, native_w: int, native_h: int) -> tuple[int | None, int | None]:
                if orig_cx is None and orig_cy is None:
                    return None, None
                if orig_cx is not None and orig_cy is not None:
                    width = orig_cx
                    height = round(width * native_h / native_w)
                    if height > orig_cy:
                        height = orig_cy
                        width = round(height * native_w / native_h)
                    return width, height
                if orig_cx is not None:
                    return orig_cx, None
                return None, orig_cy

            def _insert_picture_preserve_aspect(run, width_emus: int | None = None, height_emus: int | None = None) -> bool:
                try:
                    if width_emus is not None and height_emus is not None:
                        run.add_picture(str(image_path), width=Emu(width_emus), height=Emu(height_emus))
                    elif width_emus is not None:
                        run.add_picture(str(image_path), width=Emu(width_emus))
                    elif height_emus is not None:
                        run.add_picture(str(image_path), height=Emu(height_emus))
                    else:
                        run.add_picture(str(image_path), width=Inches(1.5))

                    inserted_cx, inserted_cy = _get_run_image_extent(run)
                    if inserted_cx is None and inserted_cy is None:
                        return False

                    target_width = inserted_cx or width_emus
                    target_height = inserted_cy or height_emus
                    if target_width is None and target_height is None:
                        return False
                    if target_width is None:
                        target_width = round(target_height * 1.0)
                    if target_height is None:
                        target_height = round(target_width * 1.0)

                    _set_extent_on_run(run, target_width, target_height)
                    return True
                except Exception:
                    _remove_drawings_from_run(run)
                    return False

            inserted = False

            def _process_paragraph_for_logo(paragraph) -> bool:
                for run in list(paragraph.runs):
                    orig_cx, orig_cy = _get_run_image_extent(run)
                    if orig_cx is None and orig_cy is None:
                        continue
                    _remove_drawings_from_run(run)
                    native_w, native_h = _get_image_native_size(image_path)
                    if desired_width is not None or desired_height is not None:
                        box_width, box_height = _calculate_fit_size(desired_width, desired_height, native_w, native_h)
                        new_run = paragraph.add_run()
                        if _insert_picture_preserve_aspect(new_run, width_emus=box_width):
                            return True
                    target_width, target_height = _calculate_fit_size(orig_cx, orig_cy, native_w, native_h)
                    new_run = paragraph.add_run()
                    if target_width is not None:
                        if _insert_picture_preserve_aspect(new_run, width_emus=target_width):
                            return True
                    else:
                        if _insert_picture_preserve_aspect(new_run, height_emus=target_height):
                            return True
                if "{{Logo}}" in paragraph.text:
                    paragraph.text = ""
                    run = paragraph.add_run()
                    native_w, native_h = _get_image_native_size(image_path)
                    if desired_width is not None or desired_height is not None:
                        box_width, box_height = _calculate_fit_size(desired_width, desired_height, native_w, native_h)
                        _insert_picture_preserve_aspect(run, width_emus=box_width)
                    else:
                        target_width, target_height = _calculate_fit_size(int(1.5 * EMU_PER_INCH), None, native_w, native_h)
                        _insert_picture_preserve_aspect(run, width_emus=target_width, height_emus=target_height)
                    return True
                return False

            def _find_extent_in_element(elem) -> int | None:
                try:
                    for el in elem.iter():
                        tag = getattr(el, 'tag', '')
                        if tag and tag.endswith('}extent'):
                            cx = el.get('cx')
                            if cx:
                                return int(cx)
                except Exception:
                    return None
                return None

            native_w, native_h = _get_image_native_size(image_path)

            for section in document.sections:
                header = section.header
                if header is None:
                    continue
                for paragraph in header.paragraphs:
                    if _process_paragraph_for_logo(paragraph):
                        inserted = True
                        break
                if inserted:
                    break
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_target = _find_extent_in_element(cell._tc)
                            for paragraph in cell.paragraphs:
                                if cell_target is not None:
                                    for run in list(paragraph.runs):
                                        orig_cx, orig_cy = _get_run_image_extent(run)
                                        if orig_cx is None and orig_cy is None:
                                            continue
                                        _remove_drawings_from_run(run)
                                        target_width, target_height = _calculate_fit_size(orig_cx, orig_cy, native_w, native_h)
                                        if target_width is not None:
                                            if _insert_picture_preserve_aspect(run, width_emus=target_width):
                                                inserted = True
                                                break
                                        else:
                                            if _insert_picture_preserve_aspect(run, height_emus=target_height):
                                                inserted = True
                                                break
                                    if inserted:
                                        break
                                else:
                                    if _process_paragraph_for_logo(paragraph):
                                        inserted = True
                                        break
                            if inserted:
                                break
                        if inserted:
                            break
                    if inserted:
                        break
                if inserted:
                    break

            if not inserted:
                for paragraph in document.paragraphs:
                    for run in list(paragraph.runs):
                        orig_cx, orig_cy = _get_run_image_extent(run)
                        if orig_cx is None and orig_cy is None:
                            continue
                        _remove_drawings_from_run(run)
                        target_width, target_height = _calculate_fit_size(orig_cx, orig_cy, native_w, native_h)
                        if target_width is not None:
                            if _insert_picture_preserve_aspect(run, width_emus=target_width):
                                inserted = True
                                break
                        else:
                            if _insert_picture_preserve_aspect(run, height_emus=target_height):
                                inserted = True
                                break
                    if inserted:
                        break
                    if "{{Logo}}" in paragraph.text:
                        paragraph.text = ""
                        run = paragraph.add_run()
                        target_width, target_height = _calculate_fit_size(int(1.5 * EMU_PER_INCH), None, native_w, native_h)
                        _insert_picture_preserve_aspect(run, width_emus=target_width, height_emus=target_height)
                        inserted = True
                        break

            if not inserted:
                if document.sections:
                    first_section = document.sections[0]
                    header = first_section.header
                    try:
                        target_width = None
                        if header is not None:
                            for paragraph in header.paragraphs:
                                for run in list(paragraph.runs):
                                    orig_cx, orig_cy = _get_run_image_extent(run)
                                    if orig_cx:
                                        target_width = orig_cx
                                    _remove_drawings_from_run(run)
                        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                        run = p.add_run()
                        if desired_width is not None or desired_height is not None:
                            box_width, box_height = _calculate_fit_size(desired_width, desired_height, native_w, native_h)
                            _insert_picture_preserve_aspect(run, width_emus=box_width)
                        else:
                            target_width, target_height = _calculate_fit_size(target_width, None, native_w, native_h)
                            _insert_picture_preserve_aspect(run, width_emus=target_width, height_emus=target_height)
                        self.logger.debug("Inserted logo image into header")
                        inserted = True
                    except Exception:
                        document.add_picture(str(image_path), width=Inches(1.5))
                        self.logger.debug("Inserted logo image at document end (fallback)")
                else:
                    document.add_picture(str(image_path), width=Inches(1.5))
                    self.logger.debug("Inserted logo image at document end")
        except Exception as error:
            self.logger.exception("Failed to insert logo image: %s", error)

    def _build_destination(
        self,
        source_path: Path,
        output_root: Path,
        preserve_folder_structure: bool,
        input_root: Path | None,
    ) -> Path:
        if preserve_folder_structure and input_root is not None:
            relative = source_path.relative_to(input_root)
            return output_root / relative
        return output_root / source_path.name
