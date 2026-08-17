from __future__ import annotations

import json
import logging
import re
import sqlite3
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class BrandingSettings:
    """Configuration for company-name replacement in document branding."""

    current_company_name: str = ""
    new_company_name: str = ""
    replace_in_header: bool = True
    replace_in_footer: bool = False
    replace_in_body: bool = False
    replace_in_tables: bool = False
    replace_in_text_boxes: bool = False


@dataclass(slots=True)
class BrandingResult:
    """Outcome of processing a single document for branding."""

    document_name: str
    header_found: bool = False
    company_name_detected: str | None = None
    company_name_replaced: bool = False
    replacements_made: int = 0
    footer_replacements: int = 0
    body_replacements: int = 0
    table_replacements: int = 0
    processing_time_seconds: float = 0.0
    success: bool = False
    error: str | None = None


class BrandingProcessor:
    """Detect and replace branding text such as company names across Word document parts."""

    def __init__(self, learning_db_path: str | Path | None = None) -> None:
        self.logger = logger
        self.learning_db_path = Path(learning_db_path or Path("config") / "branding_learning.sqlite")
        self.learning_db_path.parent.mkdir(parents=True, exist_ok=True)
        self._init_db()

    def _init_db(self) -> None:
        with sqlite3.connect(self.learning_db_path) as connection:
            connection.execute(
                """
                CREATE TABLE IF NOT EXISTS branding_learning (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    old_value TEXT NOT NULL,
                    new_value TEXT NOT NULL,
                    context TEXT NOT NULL,
                    created_at TEXT NOT NULL
                )
                """
            )
            connection.commit()

    def apply_to_document(self, document_path: Path | str, settings: BrandingSettings) -> BrandingResult:
        """Apply branding settings to a Word document and return a result summary."""
        started_at = time.perf_counter()
        path = Path(document_path)
        result = BrandingResult(document_name=path.name)
        try:
            document = Document(path)
            result = self.apply_to_document_object(document, settings, document_name=path.name)
            document.save(path)
            result.processing_time_seconds = time.perf_counter() - started_at
            return result
        except Exception as error:  # pragma: no cover - defensive path
            result.error = str(error)
            result.success = False
            result.processing_time_seconds = time.perf_counter() - started_at
            self.logger.exception("Failed to apply branding to %s", path)
            return result

    def apply_to_document_object(
        self,
        document: DocxDocument,
        settings: BrandingSettings,
        document_name: str | None = None,
    ) -> BrandingResult:
        """Apply branding settings directly to an in-memory Word document."""
        started_at = time.perf_counter()
        result = BrandingResult(document_name=document_name or "document")
        try:
            target_name = self._resolve_company_name(settings)
            if not target_name:
                result.error = "No replacement company name supplied"
                result.success = False
                result.processing_time_seconds = time.perf_counter() - started_at
                return result

            replacements = 0
            result.company_name_detected = self._detect_company_name(document)
            if not result.company_name_detected and not settings.current_company_name:
                self.logger.warning("Could not detect current company name for %s", document_name or "document")

            if settings.replace_in_header:
                replacements += self._replace_in_sections(document, result.company_name_detected or settings.current_company_name, target_name, section_type="header")
                result.header_found = True
            if settings.replace_in_footer:
                replacements += self._replace_in_sections(document, result.company_name_detected or settings.current_company_name, target_name, section_type="footer")
            if settings.replace_in_body:
                replacements += self._replace_in_body(document, result.company_name_detected or settings.current_company_name, target_name)
            if settings.replace_in_tables:
                replacements += self._replace_in_tables(document, result.company_name_detected or settings.current_company_name, target_name)
            if settings.replace_in_text_boxes:
                replacements += self._replace_in_text_boxes(document, result.company_name_detected or settings.current_company_name, target_name)

            result.replacements_made = replacements
            result.company_name_replaced = replacements > 0
            result.success = True
            result.processing_time_seconds = time.perf_counter() - started_at
            return result
        except Exception as error:  # pragma: no cover - defensive path
            result.error = str(error)
            result.success = False
            result.processing_time_seconds = time.perf_counter() - started_at
            self.logger.exception("Failed to apply branding to %s", document_name or "document")
            return result

    def _resolve_company_name(self, settings: BrandingSettings) -> str:
        if settings.new_company_name:
            return settings.new_company_name.strip()
        return ""

    def _detect_company_name(self, document: DocxDocument) -> str | None:
        candidates: list[str] = []
        for section in document.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header is None:
                    continue
                for paragraph in header.paragraphs:
                    text = self._normalize_text(paragraph.text)
                    if text:
                        candidates.append(text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = self._normalize_text(paragraph.text)
                                if text:
                                    candidates.append(text)
        if not candidates:
            return None
        return max(candidates, key=len)

    def _replace_in_sections(self, document: DocxDocument, old_name: str | None, new_name: str, section_type: str) -> int:
        if not old_name:
            return 0
        replacements = 0
        for section in document.sections:
            section_items: list[Any] = []
            if section_type == "header":
                section_items = [section.header, section.first_page_header, section.even_page_header]
            elif section_type == "footer":
                section_items = [section.footer, section.first_page_footer, section.even_page_footer]

            for section_part in section_items:
                if section_part is None:
                    continue
                for paragraph in section_part.paragraphs:
                    replacements += self._replace_text_in_paragraph(paragraph, old_name, new_name)
                for table in section_part.tables:
                    replacements += self._replace_text_in_table(table, old_name, new_name)
        return replacements

    def _replace_in_body(self, document: DocxDocument, old_name: str | None, new_name: str) -> int:
        if not old_name:
            return 0
        replacements = 0
        for paragraph in document.paragraphs:
            replacements += self._replace_text_in_paragraph(paragraph, old_name, new_name)
        return replacements

    def _replace_in_tables(self, document: DocxDocument, old_name: str | None, new_name: str) -> int:
        if not old_name:
            return 0
        replacements = 0
        for table in document.tables:
            replacements += self._replace_text_in_table(table, old_name, new_name)
        return replacements

    def _replace_in_text_boxes(self, document: DocxDocument, old_name: str | None, new_name: str) -> int:
        if not old_name:
            return 0
        return 0

    def _replace_text_in_paragraph(self, paragraph: Paragraph, old_name: str, new_name: str) -> int:
        if not old_name:
            return 0
        text = paragraph.text
        if not text:
            return 0

        normalized_old = self._normalize_text(old_name)
        normalized_text = self._normalize_text(text)
        if not normalized_old or normalized_old not in normalized_text:
            return 0

        updated_text = re.sub(re.escape(old_name), lambda _match: new_name, text, flags=re.IGNORECASE)
        if updated_text != text:
            paragraph.text = updated_text
            return 1

        for run in paragraph.runs:
            run_text = run.text or ""
            if not run_text:
                continue
            if self._normalize_text(run_text) and normalized_old in self._normalize_text(run_text):
                run.text = re.sub(re.escape(old_name), lambda _match: new_name, run_text, flags=re.IGNORECASE)
                return 1

        return 0

    def _replace_text_in_table(self, table: Table, old_name: str, new_name: str) -> int:
        replacements = 0
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replacements += self._replace_text_in_paragraph(paragraph, old_name, new_name)
        return replacements

    def _normalize_text(self, value: str) -> str:
        cleaned = re.sub(r"\s+", " ", (value or "")).strip().lower()
        return cleaned

    def learn(self, old_value: str, new_value: str, context: str = "branding") -> None:
        """Store a branding replacement example for future reuse."""
        if not old_value or not new_value:
            return
        with sqlite3.connect(self.learning_db_path) as connection:
            connection.execute(
                "INSERT INTO branding_learning (old_value, new_value, context, created_at) VALUES (?, ?, ?, ?)",
                (old_value.strip(), new_value.strip(), context, time.strftime("%Y-%m-%d %H:%M:%S")),
            )
            connection.commit()

    def suggest_replacement(self, old_value: str, context: str = "branding") -> str | None:
        """Suggest a previous replacement for a known old company name."""
        if not old_value:
            return None
        with sqlite3.connect(self.learning_db_path) as connection:
            row = connection.execute(
                "SELECT new_value FROM branding_learning WHERE lower(old_value)=? AND lower(context)=? ORDER BY id DESC LIMIT 1",
                (old_value.strip().lower(), context.lower()),
            ).fetchone()
        return row[0] if row else None
