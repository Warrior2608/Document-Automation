from __future__ import annotations

import json
import logging
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class MetadataValues:
    """Values collected from the metadata replacement form."""

    effective_date: str = ""
    latest_revision_date: str = ""
    next_review_date: str = ""
    issue: str = ""
    version: str = ""
    prepared_by_name: str = ""
    prepared_by_designation: str = ""
    checked_by_name: str = ""
    checked_by_designation: str = ""
    approved_by_name: str = ""
    approved_by_designation: str = ""

    def to_dict(self) -> dict[str, str]:
        return {
            "effective_date": self.effective_date,
            "latest_revision_date": self.latest_revision_date,
            "next_review_date": self.next_review_date,
            "issue": self.issue,
            "version": self.version,
            "prepared_by_name": self.prepared_by_name,
            "prepared_by_designation": self.prepared_by_designation,
            "checked_by_name": self.checked_by_name,
            "checked_by_designation": self.checked_by_designation,
            "approved_by_name": self.approved_by_name,
            "approved_by_designation": self.approved_by_designation,
        }


@dataclass(slots=True)
class MetadataUpdateResult:
    """Outcome of processing a single Word document."""

    document_name: str
    metadata_table_found: bool
    updated_fields: list[str] = field(default_factory=list)
    missing_fields: list[str] = field(default_factory=list)
    processing_time_seconds: float = 0.0
    success: bool = False
    error: str | None = None


class MetadataProcessor:
    """Replace metadata values in the first-page metadata table of a Word document."""

    LABEL_MAP: dict[str, str] = {
        "effective date": "effective_date",
        "latest revision date": "latest_revision_date",
        "next review date": "next_review_date",
        "issue": "issue",
        "version": "version",
        "prepared by": "prepared_by",
        "checked by": "checked_by",
        "approved by": "approved_by",
    }

    def __init__(self) -> None:
        self.logger = logger

    def apply_to_document(self, document_path: Path | str, values: MetadataValues) -> MetadataUpdateResult:
        """Apply metadata values to a Word file and return detailed processing results."""
        started_at = time.perf_counter()
        path = Path(document_path)
        result = MetadataUpdateResult(document_name=path.name, metadata_table_found=False)
        try:
            document = Document(path)
            result = self.apply_to_document_object(document, values, document_name=path.name)
            document.save(path)
            result.processing_time_seconds = time.perf_counter() - started_at
            return result
        except Exception as error:  # pragma: no cover - defensive path
            result.error = str(error)
            result.success = False
            result.processing_time_seconds = time.perf_counter() - started_at
            self.logger.exception("Failed to process metadata for %s", path)
            return result

    def apply_to_document_object(
        self,
        document: DocxDocument,
        values: MetadataValues,
        document_name: str | None = None,
    ) -> MetadataUpdateResult:
        """Apply metadata values directly to an in-memory Word document."""
        started_at = time.perf_counter()
        result = MetadataUpdateResult(document_name=document_name or "document", metadata_table_found=False)
        try:
            table = self._find_metadata_table(document)
            if table is None:
                result.error = "Metadata table not found"
                result.success = False
                self.logger.warning("Metadata table not found in %s", document_name or "document")
                result.processing_time_seconds = time.perf_counter() - started_at
                return result

            result.metadata_table_found = True
            values_by_field = self._build_values(values)
            for label, key in self.LABEL_MAP.items():
                field_key = self._field_key_from_label(label)
                if field_key is None:
                    continue
                value = values_by_field.get(field_key)
                if value:
                    self._replace_label_value(table, label, value)
                    result.updated_fields.append(field_key)
                else:
                    result.missing_fields.append(field_key)

            result.success = True
            result.processing_time_seconds = time.perf_counter() - started_at
            return result
        except Exception as error:  # pragma: no cover - defensive path
            result.error = str(error)
            result.success = False
            result.processing_time_seconds = time.perf_counter() - started_at
            self.logger.exception("Failed to process metadata for %s", document_name or "document")
            return result

    def apply_to_documents(self, documents: list[Path], values: MetadataValues) -> list[MetadataUpdateResult]:
        """Process a batch of documents and collect per-document results."""
        return [self.apply_to_document(path, values) for path in documents]

    def save_defaults(self, values: MetadataValues, config_path: Path | str | None = None) -> Path:
        """Persist metadata values to a local JSON file."""
        path = Path(config_path or Path("config") / "metadata_defaults.json")
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(values.to_dict(), indent=2), encoding="utf-8")
        return path

    def load_defaults(self, config_path: Path | str | None = None) -> MetadataValues:
        """Load saved metadata values from a local JSON file."""
        path = Path(config_path or Path("config") / "metadata_defaults.json")
        if not path.exists():
            return MetadataValues()
        payload = json.loads(path.read_text(encoding="utf-8"))
        return MetadataValues(
            effective_date=payload.get("effective_date", ""),
            latest_revision_date=payload.get("latest_revision_date", ""),
            next_review_date=payload.get("next_review_date", ""),
            issue=payload.get("issue", ""),
            version=payload.get("version", ""),
            prepared_by_name=payload.get("prepared_by_name", ""),
            prepared_by_designation=payload.get("prepared_by_designation", ""),
            checked_by_name=payload.get("checked_by_name", ""),
            checked_by_designation=payload.get("checked_by_designation", ""),
            approved_by_name=payload.get("approved_by_name", ""),
            approved_by_designation=payload.get("approved_by_designation", ""),
        )

    def _build_values(self, values: MetadataValues) -> dict[str, str]:
        return {
            "effective_date": values.effective_date,
            "latest_revision_date": values.latest_revision_date,
            "next_review_date": values.next_review_date,
            "issue": values.issue,
            "version": values.version,
            "prepared_by": self._format_name_designation(values.prepared_by_name, values.prepared_by_designation),
            "checked_by": self._format_name_designation(values.checked_by_name, values.checked_by_designation),
            "approved_by": self._format_name_designation(values.approved_by_name, values.approved_by_designation),
        }

    def _format_name_designation(self, name: str, designation: str) -> str:
        name = name.strip()
        designation = designation.strip()
        if not name and not designation:
            return ""
        if not name:
            return designation
        if not designation:
            return name
        return f"{name} / {designation}"

    def _find_metadata_table(self, document: DocxDocument) -> Table | None:
        """Find the first page metadata table by scanning document tables and rows."""
        for table in document.tables:
            if self._looks_like_metadata_table(table):
                return table
        return None

    def _looks_like_metadata_table(self, table: Table) -> bool:
        """Heuristically identify a metadata table containing label/value pairs."""
        for row in table.rows[:8]:
            if len(row.cells) < 2:
                continue
            first_text = self._normalize_label(row.cells[0].text)
            if first_text in {
                "effective date",
                "latest revision date",
                "next review date",
                "issue",
                "version",
                "prepared by",
                "checked by",
                "approved by",
            }:
                return True
        return False

    def _replace_label_value(self, table: Table, label: str, value: str) -> None:
        normalized_label = self._normalize_label(label)
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label_text = self._normalize_label(row.cells[0].text)
            if self._matches_label(label_text, normalized_label):
                self._set_cell_text(row.cells[1], value)
                return

    def _field_key_from_label(self, label: str) -> str | None:
        normalized_label = self._normalize_label(label)
        if normalized_label.startswith("effective date"):
            return "effective_date"
        if normalized_label.startswith("latest revision date"):
            return "latest_revision_date"
        if normalized_label.startswith("next review date"):
            return "next_review_date"
        if normalized_label.startswith("issue"):
            return "issue"
        if normalized_label.startswith("version"):
            return "version"
        if normalized_label.startswith("prepared by"):
            return "prepared_by"
        if normalized_label.startswith("checked by"):
            return "checked_by"
        if normalized_label.startswith("approved by"):
            return "approved_by"
        return None

    def _matches_label(self, row_label: str, target_label: str) -> bool:
        if row_label == target_label:
            return True
        if target_label.startswith("prepared by") and row_label.startswith("prepared by"):
            return True
        if target_label.startswith("checked by") and row_label.startswith("checked by"):
            return True
        if target_label.startswith("approved by") and row_label.startswith("approved by"):
            return True
        return False

    def _set_cell_text(self, cell, text: str) -> None:
        if cell.paragraphs:
            first_paragraph = cell.paragraphs[0]
            first_paragraph.text = text
            for paragraph in cell.paragraphs[1:]:
                paragraph.clear()
            return
        paragraph = cell.add_paragraph()
        paragraph.text = text

    def _normalize_label(self, text: str) -> str:
        cleaned = re.sub(r"\s+", " ", text or "").strip().lower().rstrip(":")
        return cleaned
