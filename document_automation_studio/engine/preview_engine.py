from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from openpyxl import load_workbook

from document_automation_studio.engine.rule_engine import RuleEngine
from document_automation_studio.models.preview_models import PreviewChange
from document_automation_studio.models.rule_models import RuleSet, TextReplacementRule

logger = logging.getLogger(__name__)


class PreviewEngine:
    """Generate preview items for document replacements and rename suggestions."""

    def __init__(self, rule_engine: RuleEngine) -> None:
        self.rule_engine = rule_engine

    def preview_files(
        self,
        files: Iterable[Path],
        rule_set: RuleSet | None = None,
        input_root: Path | None = None,
        max_items: int = 20,
        process_word_files: bool = True,
        process_excel_files: bool = True,
        logo_path: str | None = None,
    ) -> list[PreviewChange]:
        changes: list[PreviewChange] = []
        for file_path in files:
            if len(changes) >= max_items:
                break
            suffix = file_path.suffix.lower()
            if suffix == ".docx" and process_word_files:
                changes.extend(self._preview_word_file(file_path, rule_set or RuleSet(), input_root, logo_path=logo_path))
            elif suffix == ".xlsx" and process_excel_files:
                changes.extend(self._preview_excel_file(file_path, rule_set or RuleSet(), input_root))

        return changes[:max_items]

    def _preview_word_file(self, source_path: Path, rule_set: RuleSet, input_root: Path | None, logo_path: str | None = None) -> list[PreviewChange]:
        preview_items: list[PreviewChange] = []
        text_replacements = rule_set.text_replacements if rule_set else []

        document = Document(source_path)
        for paragraph in document.paragraphs:
            preview_items.extend(self._preview_paragraph(source_path, paragraph, text_replacements))
        for table in document.tables:
            preview_items.extend(self._preview_table(source_path, table, text_replacements))
        for section in document.sections:
            if section.header is not None:
                for paragraph in section.header.paragraphs:
                    preview_items.extend(self._preview_paragraph(source_path, paragraph, text_replacements))
            if section.footer is not None:
                for paragraph in section.footer.paragraphs:
                    preview_items.extend(self._preview_paragraph(source_path, paragraph, text_replacements))

        if logo_path:
            preview_items.extend(self._preview_logo(source_path, logo_path))
        if rule_set and rule_set.rename_patterns:
            preview_items.extend(self._preview_rename(source_path, rule_set, input_root))

        return preview_items

    def _preview_excel_file(self, source_path: Path, rule_set: RuleSet, input_root: Path | None) -> list[PreviewChange]:
        preview_items: list[PreviewChange] = []
        replacements = rule_set.excel_replacements if rule_set.excel_replacements else rule_set.text_replacements
        if replacements:
            workbook = load_workbook(source_path, data_only=True)
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=False):
                    for cell in row:
                        if isinstance(cell.value, str):
                            original = cell.value
                            new_value = self.rule_engine.apply_text_replacements(original, replacements)
                            if original != new_value:
                                preview_items.append(
                                    PreviewChange(
                                        file_path=source_path,
                                        change_type="Excel Replacement",
                                        old_value=original,
                                        new_value=new_value,
                                        details=f"{sheet.title}:{cell.coordinate}",
                                    )
                                )

        if rule_set.excel_insert_rows:
            preview_items.extend(self._preview_insert_rows(source_path, rule_set.excel_insert_rows))
        if rule_set.excel_insert_columns:
            preview_items.extend(self._preview_insert_columns(source_path, rule_set.excel_insert_columns))
        if rule_set.excel_hyperlinks:
            preview_items.extend(self._preview_hyperlinks(source_path, rule_set.excel_hyperlinks))
        if rule_set.rename_patterns:
            preview_items.extend(self._preview_rename(source_path, rule_set, input_root))

        return preview_items

    def _preview_insert_rows(self, source_path: Path, row_rules: list[dict[str, object]]) -> list[PreviewChange]:
        preview_items: list[PreviewChange] = []
        for rule in row_rules:
            sheet_name = rule.get("sheet") or "active"
            index = int(rule.get("index", 1))
            values = rule.get("values", [])
            preview_items.append(
                PreviewChange(
                    file_path=source_path,
                    change_type="Excel Insert Row",
                    old_value="",
                    new_value=f"Insert row at {sheet_name}:{index} values={values}",
                    details="Insert row rule",
                )
            )
        return preview_items

    def _preview_insert_columns(self, source_path: Path, column_rules: list[dict[str, object]]) -> list[PreviewChange]:
        preview_items: list[PreviewChange] = []
        for rule in column_rules:
            sheet_name = rule.get("sheet") or "active"
            index = int(rule.get("index", 1))
            values = rule.get("values", [])
            preview_items.append(
                PreviewChange(
                    file_path=source_path,
                    change_type="Excel Insert Column",
                    old_value="",
                    new_value=f"Insert column at {sheet_name}:{index} values={values}",
                    details="Insert column rule",
                )
            )
        return preview_items

    def _preview_hyperlinks(self, source_path: Path, hyperlink_rules: list[dict[str, str]]) -> list[PreviewChange]:
        preview_items: list[PreviewChange] = []
        for rule in hyperlink_rules:
            sheet_name = rule.get("sheet") or "active"
            target = rule.get("cell", "")
            url = rule.get("url", "")
            display = rule.get("display", url)
            preview_items.append(
                PreviewChange(
                    file_path=source_path,
                    change_type="Excel Hyperlink",
                    old_value="",
                    new_value=f"Set hyperlink {sheet_name}:{target} -> {display} ({url})",
                    details="Hyperlink rule",
                )
            )
        return preview_items

    def _preview_paragraph(self, source_path: Path, paragraph: Paragraph, replacements: Iterable[TextReplacementRule]) -> list[PreviewChange]:
        changes: list[PreviewChange] = []
        original = paragraph.text
        new_text = self.rule_engine.apply_text_replacements(original, replacements)
        if original != new_text:
            changes.append(
                PreviewChange(
                    file_path=source_path,
                    change_type="Text Replacement",
                    old_value=original,
                    new_value=new_text,
                    details="Paragraph text",
                )
            )
        return changes

    def _preview_table(self, source_path: Path, table: _Cell | object, replacements: Iterable[TextReplacementRule]) -> list[PreviewChange]:
        changes: list[PreviewChange] = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changes.extend(self._preview_paragraph(source_path, paragraph, replacements))
        return changes

    def _preview_logo(self, source_path: Path, logo_path: str | None) -> list[PreviewChange]:
        preview_items: list[PreviewChange] = []
        if not logo_path:
            return preview_items

        document = Document(source_path)
        for paragraph in document.paragraphs:
            if "{{Logo}}" in paragraph.text:
                preview_items.append(
                    PreviewChange(
                        file_path=source_path,
                        change_type="Logo Insertion",
                        old_value="{{Logo}}",
                        new_value=f"Insert logo from {logo_path}",
                        details="Logo placeholder detected",
                    )
                )
                return preview_items

        preview_items.append(
            PreviewChange(
                file_path=source_path,
                change_type="Logo Insertion",
                old_value="",
                new_value=f"Insert logo from {logo_path} at document end",
                details="No placeholder found",
            )
        )
        return preview_items

    def _preview_rename(self, source_path: Path, rule_set: RuleSet, input_root: Path | None = None) -> list[PreviewChange]:
        changes: list[PreviewChange] = []
        if not rule_set.rename_patterns:
            return changes

        old_name = source_path.name
        new_name = old_name
        for find_text, replace_text in rule_set.rename_patterns.items():
            new_name = new_name.replace(find_text, replace_text)

        if new_name != old_name:
            changes.append(
                PreviewChange(
                    file_path=source_path,
                    change_type="File Rename",
                    old_value=old_name,
                    new_value=new_name,
                    details="Rename pattern",
                )
            )
        return changes
