from pathlib import Path
from document_automation_studio.processors.word_processor import WordProcessor
from docx import Document

# Prepare paths
base = Path(__file__).resolve().parent
doc_path = base / "logo_test.docx"
image_path = base / "test_logo.png"
output_dir = base / "out"
output_dir.mkdir(exist_ok=True)

# Create a tiny image if not exists
if not image_path.exists():
    from PIL import Image, ImageDraw
    img = Image.new('RGBA', (200, 100), (255, 0, 0, 255))
    d = ImageDraw.Draw(img)
    d.text((10,40), "LOGO", fill=(255,255,255))
    img.save(image_path)

# Create a docx with header placeholder
doc = Document()
section = doc.sections[0]
header = section.header
# create a 1x2 table in header and insert the image into the left cell
from docx.shared import Inches as _Inches
table = header.add_table(rows=1, cols=2, width=_Inches(6))
cell = table.cell(0, 0)
cell_p = cell.paragraphs[0]
run = cell_p.add_run()
run.add_picture(str(image_path), width=_Inches(2.0))
doc.add_paragraph("Body text here")
orig_path = base / "logo_test_original.docx"
doc.save(orig_path)

# Run processor to replace header image with new image (smaller)
processor = WordProcessor()
out = processor.process(orig_path, output_dir, preserve_folder_structure=False, replacements=None, rule_set=None, logo_path=str(image_path))
print('Saved:', out)
print('Exists:', out.exists())

# Inspect the output to read back header image size (EMU -> inches)
from docx.oxml.ns import qn
out_doc = Document(out)
found_width_inches = None
ns = {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}
# Debug: inspect original header extents
print('\n-- Original header extents --')
orig_doc = Document(orig_path)
for section in orig_doc.sections:
    header = section.header
    if not header:
        continue
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            for el in run._element.iter():
                tag = getattr(el, 'tag', '')
                if tag and tag.endswith('}extent'):
                    print('Found extent in run:', el.get('cx'))
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                for el in cell._tc.iter():
                    tag = getattr(el, 'tag', '')
                    if tag and tag.endswith('}extent'):
                        print('Found extent in cell:', el.get('cx'))
print('-- end orig extents --\n')
# Debug run XML for orig header cell
print('\n-- Original run XML --')
for section in orig_doc.sections:
    header = section.header
    if not header:
        continue
    for paragraph in header.paragraphs:
        for idx, run in enumerate(paragraph.runs):
            print('orig paragraph run', idx, run._element.xml)
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                for idx, paragraph in enumerate(cell.paragraphs):
                    for ridx, run in enumerate(paragraph.runs):
                        print('orig cell paragraph', idx, 'run', ridx, run._element.xml)
print('-- end original run XML --\n')

# Also list extents found in the output header
print('\n-- Output header extents --')
for section in out_doc.sections:
    header = section.header
    if not header:
        continue
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            for el in run._element.iter():
                tag = getattr(el, 'tag', '')
                if tag and tag.endswith('}extent'):
                    print('Found extent in run (out):', el.get('cx'))
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                for el in cell._tc.iter():
                    tag = getattr(el, 'tag', '')
                    if tag and tag.endswith('}extent'):
                        print('Found extent in cell (out):', el.get('cx'))
print('-- end out extents --\n')

print('\n-- Output run XML --')
for section in out_doc.sections:
    header = section.header
    if not header:
        continue
    for paragraph in header.paragraphs:
        for idx, run in enumerate(paragraph.runs):
            print('out paragraph run', idx, run._element.xml)
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                for idx, paragraph in enumerate(cell.paragraphs):
                    for ridx, run in enumerate(paragraph.runs):
                        print('out cell paragraph', idx, 'run', ridx, run._element.xml)
print('-- end output run XML --\n')

# Try to extract a representative width
for section in out_doc.sections:
    header = section.header
    if header is None:
        continue
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            for el in run._element.iter():
                tag = getattr(el, 'tag', '')
                if tag and tag.endswith('}extent'):
                    cx = el.get('cx')
                    if cx:
                        found_width_inches = int(cx) / 914400
                        break
            if found_width_inches:
                break
        if found_width_inches:
            break
    if found_width_inches:
        break
print('Replaced header image width (inches):', found_width_inches)
